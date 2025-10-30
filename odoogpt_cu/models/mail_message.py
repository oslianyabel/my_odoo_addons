import base64
import logging
import os
import secrets
from datetime import datetime
from io import BytesIO

import markdown2
import openpyxl
from docx import Document
from markupsafe import Markup
from odoo import _, api, exceptions, models  # type: ignore
from PyPDF2 import PdfReader

from .completions import agent
from .enumerations import MessageType
from .prompt import JSON_TOOLS, SYSTEM_PROMPT
from .tools import send_odoo_msg, tools_func
from .utils import UserPhone, format_phone_number

_logger = logging.getLogger(__name__)
ENV = os.getenv("ENV", "prod")


class MailMessage(models.Model):
    _inherit = "mail.message"

    @api.model_create_multi
    def create(self, values_list):
        # Override to create a new job for all created messages in the odoogpt channel
        res = super(MailMessage, self).create(values_list)
        try:
            odoogpt = self.env.ref("odoogpt.partner_odoogpt")
        except Exception:
            return res
        chats = res.filtered(
            lambda s: s.model == "discuss.channel" and s.author_id.id != odoogpt.id
        )
        mail_channel_obj = self.env["discuss.channel"].sudo()
        for chat in chats:
            channel_id = mail_channel_obj.browse(chat.res_id).exists()
            if channel_id.is_odoogpt_chat:
                if ENV == "dev":
                    chat._send_message_to_odoogpt(odoogpt, channel_id)
                else:
                    chat.with_delay()._send_message_to_odoogpt(odoogpt, channel_id)

        return res

    def _send_message_to_odoogpt(self, odoogpt, channel_id):
        # Return the response from openai api and send a message
        self.ensure_one()

        ai_msg = self._gen_ai_ans(odoogpt, channel_id)
        if ai_msg:
            channel_id.message_post(
                body=Markup(ai_msg),
                message_type="comment",
                subtype_xmlid="mail.mt_comment",
                author_id=odoogpt.id,
                email_from=odoogpt.email,
            )

    def _get_table_data(self, table_name):
        try:
            self.ensure_one()
            if "_" in table_name:
                table_name = table_name.replace("_", ".")

            if table_name not in self.env:
                raise ValueError(f"No table named {table_name} exists.")

            table_records = self.env[table_name].sudo().search([])
            if not table_records:
                raise ValueError(f"No value in {table_name} table.")

            table_info = {
                record.id: {
                    field: getattr(record, field)
                    for field in self.env[table_name]._fields.keys()
                }
                for record in table_records
            }
            return str(table_info)

        except Exception:
            raise ValueError(f"Problem with the {table_name} table.")

    def _gen_ai_ans(self, odoogpt, channel_id) -> str:
        """
        Send a question to the openai api using per-chat Completions instance.
        :param odoogpt: partner odoogpt to build a memory
        :param channel_id: discuss.channel object to build a memory
        :return: str: openai response
        """
        self.ensure_one()

        # Prepare current date context
        today = datetime.now().strftime("%Y-%m-%d")
        date_context = f"Fecha actual: {today}"

        # Get only the current user message (no conversation history)
        human_input = str(self.body).replace("<p>", "").replace("</p>", "")

        # Extract context from supported attachments and prepend as system context
        attachment_context = self._build_attachments_context()

        if attachment_context:
            agent.chat_memory.add_msg(
                message=attachment_context,
                role=MessageType.SYSTEM.value,
                user_id=channel_id.id,
            )

        full_message = f"{date_context}\n\nUsuario: {human_input}"
        raw_response = agent.process_msg(
            message=full_message,
            user_id=channel_id.id,
            channel_obj=channel_id,
            odoo_manager=self,
            odoogpt=odoogpt,
            rag_functions=tools_func,
            rag_prompt=JSON_TOOLS,
            tool_execution_callback=send_odoo_msg,
        )
        return self._format_response_for_odoo(raw_response)

    def _format_response_for_odoo(self, markdown):
        # Convert model Markdown to rich HTML to display nicely in Discuss
        # Keep links, lists, code blocks, and tables for better readability
        extras = {
            "fenced-code-blocks": None,
            "tables": None,
            "strike": None,
            "task_list": None,
            "cuddled-lists": None,
            "break-on-newline": None,
        }
        html = markdown2.markdown(markdown, extras=extras)
        # Return HTML directly so Discuss renders the formatting
        return html.strip()

    def _get_memory(self, odoogpt, channel_id, limit=20):
        memory = ""
        messages = self.search(
            [
                ("model", "=", "discuss.channel"),
                ("res_id", "=", channel_id.id),
                ("id", "!=", self.id),
            ],
            order="id desc",
            limit=limit,
        )
        messages = messages.sorted(lambda s: s.id)
        for msj in messages:
            if msj.author_id.id == odoogpt.id:
                memory += (
                    "Assistant: "
                    + str(msj.body).replace("<p>", "").replace("</p>", "")
                    + "\n"
                )
            else:
                memory += (
                    "Human: "
                    + str(msj.body).replace("<p>", "").replace("</p>", "")
                    + "\n"
                )
        return memory

    # ------- Attachments extraction helpers -------
    def _build_attachments_context(self, max_total_chars: int = 6000) -> str:
        """Build a textual context from the supported attachments in this message.
        Limits total aggregated size to avoid overloading the model.
        """
        supported_mimes = {
            "application/pdf": "pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
        }
        context_parts = []
        total = 0
        for att in self.attachment_ids:
            ext = (att.name or "").lower()
            kind = None
            if att.mimetype in supported_mimes:
                kind = supported_mimes[att.mimetype]
            elif ext.endswith(".pdf"):
                kind = "pdf"
            elif ext.endswith(".docx"):
                kind = "docx"
            elif ext.endswith(".xlsx"):
                kind = "xlsx"

            if not kind:
                continue

            try:
                text = self._extract_attachment_text(att, kind)
            except Exception as e:
                _logger.warning(f"No se pudo extraer texto de {att.name}: {e}")
                continue

            if not text:
                continue

            # Truncate per attachment to be safe
            per_limit = 3000
            text = text.strip()
            if len(text) > per_limit:
                text = text[:per_limit] + "\n...[contenido truncado]"
                print(
                    f"Contenido del archivo truncado por exceder los {per_limit} caracteres"
                )

            header = f"\n---\nArchivo: {att.name}\nTipo: {kind}\nContenido:\n"
            chunk = header + text
            if total + len(chunk) > max_total_chars:
                remain = max_total_chars - total
                if remain <= 0:
                    break
                chunk = chunk[:remain] + "\n...[contexto de adjuntos truncado]"
            context_parts.append(chunk)
            total += len(chunk)

            if total >= max_total_chars:
                break

        return "".join(context_parts).strip()

    def _extract_attachment_text(self, attachment, kind: str) -> str:
        data = attachment.datas or False
        if not data:
            return ""
        raw = base64.b64decode(data)
        bio = BytesIO(raw)
        if kind == "pdf":
            return self._extract_pdf(bio)
        if kind == "docx":
            return self._extract_docx(bio)
        if kind == "xlsx":
            return self._extract_xlsx(bio)
        return ""

    def _extract_pdf(self, bio: BytesIO, max_pages: int = 20) -> str:
        text_parts = []
        reader = PdfReader(bio)
        for i, page in enumerate(reader.pages):
            if i >= max_pages:
                break
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            if txt:
                text_parts.append(txt)

        _logger.info(f"contenido del PDF de {enumerate(reader.pages)} paginas extraido")
        return "\n".join(text_parts)

    def _extract_docx(self, bio: BytesIO) -> str:
        doc = Document(bio)
        paras = [p.text for p in doc.paragraphs if p.text]
        _logger.info(
            f"Contenido del Docx de {enumerate(doc.paragraphs)} parrafos extraido"
        )
        return "\n".join(paras)

    def _extract_xlsx(self, bio: BytesIO, max_cells: int = 5000) -> str:
        wb = openpyxl.load_workbook(bio, data_only=True, read_only=True)
        out = []
        cells_count = 0
        for ws in wb.worksheets:
            out.append(f"[Hoja] {ws.title}")
            for row in ws.iter_rows(values_only=True):
                # Build a CSV-like line
                vals = []
                for v in row:
                    vals.append("" if v is None else str(v))
                line = ", ".join(vals).strip()
                if line:
                    out.append(line)
                    cells_count += len(vals)
                    if cells_count >= max_cells:
                        out.append("...[contenido de hoja truncado]")
                        break
            if cells_count >= max_cells:
                break

        _logger.info(f"Contenido del xlsx de {enumerate(wb.worksheets)} hojas extraido")
        return "\n".join(out)

    def _get_memory2(self, odoogpt, channel_id, limit=20):
        memory = [{"role": "system", "content": SYSTEM_PROMPT}]
        messages = self.search(
            [
                ("model", "=", "discuss.channel"),
                ("res_id", "=", channel_id.id),
                ("id", "!=", self.id),
            ],
            order="id desc",
            limit=limit,
        )
        messages = messages.sorted(lambda s: s.id)
        for msg in messages:
            msg_str = str(msg.body).replace("<p>", "").replace("</p>", "")
            if msg.author_id.id == odoogpt.id:
                memory.append({"role": "assistant", "content": msg_str})
            else:
                memory.append({"role": "user", "content": msg_str})

        return memory

    @api.model
    def create_odoo_record(self, model, args):
        try:
            new_record = self.env[model].sudo().create(args)
            return new_record
        except Exception as e:
            raise exceptions.UserError(f"Error creating record: {str(e)}")

    @api.model
    def fetch_odoo_records(
        self, model, domain, fields=None, limit=None, order=None, group_by=None
    ):
        try:
            records = self.env[model].sudo().search(domain, limit=limit, order=order)
            return records.read(fields) if fields else records
        except Exception as e:
            raise exceptions.UserError(f"Error fetching records: {str(e)}")

    def get_partner(
        self, phone=None, email=None, partner_id=None, name=None, phone_domain=[]
    ):
        if phone:
            domain = [("phone", "=", phone)]
        elif phone_domain:
            domain = phone_domain
        elif email:
            domain = [("email", "=", email)]
        elif partner_id:
            domain = [("id", "=", partner_id)]
        elif name:
            domain = [("name", "ilike", name)]
        else:
            domain = [("name", "!=", ""), ("email", "!=", ""), ("phone", "!=", "")]

        fields = [
            "id",
            "name",
            "phone",
            "email",
            "function",  # Job position
            "street",  # Address street
            "street2",  # Address street 2
            "city",  # City
            "state_id",  # State
            "zip",  # Postal code
            "country_id",  # Country
            "is_company",  # Whether it's a company or individual
            "parent_id",
        ]

        partner = self.fetch_odoo_records("res.partner", domain, fields=fields)
        if partner:
            _logger.info(f"Partners encontrados: {partner}")
            return partner  # pueden ser mas de 1

        _logger.info("No se encontraron usuarios")
        return None

    @api.model
    def get_partner_by_id(self, partner_id):
        _logger.info(f"Getting partner with id: {partner_id}")
        return self.get_partner(partner_id=int(partner_id))

    @api.model
    def get_partner_by_phone(self, phone):
        _logger.info(f"Getting partner with phone: {phone}")
        phone_domain = UserPhone(phone).get_domain()
        partner = self.get_partner(phone_domain=phone_domain)
        return partner

    @api.model
    def get_partner_by_email(self, email):
        _logger.info(f"Getting partner with email: {email}")
        return self.get_partner(email=email)

    @api.model
    def get_partner_by_name(self, name):
        _logger.info(f"Getting partner with name: {name}")
        return self.get_partner(name=name)

    @api.model
    def get_product(self, id=None, sku=None, name=None):
        domain = [("active", "=", True)]

        if sku:
            domain.append(("default_code", "=", sku))
        elif name:
            domain.append(("name", "ilike", name))
        elif id:
            domain.append(("id", "=", id))

        fields = [
            "default_code",
            "barcode",
            "name",
            "sale_ok",
            "purchase_ok",
            "detailed_type",
            "invoice_policy",
            "list_price",
            "taxes_id",
            "standard_price",
            "categ_id",
            "qty_available",
            "company_id",
        ]

        product = self.env["product.product"].sudo().search(domain, limit=1)

        if product:
            result = product.read(fields)[0]
            return result

        return None

    @api.model
    def get_all_products(self):
        products = self.env["product.product"].sudo().search([("active", "=", True)])
        templates = self.env["product.template"].sudo().search([("active", "=", True)])

        product_fields = [
            "default_code",
            "barcode",
            "name",
            "sale_ok",
            "purchase_ok",
            "detailed_type",
            "invoice_policy",
            "list_price",
            "taxes_id",
            "standard_price",
            "categ_id",
            "qty_available",
            "company_id",
        ]

        product_data = products.read(product_fields)
        template_data = templates.read(product_fields)

        _logger.info(f"Products amount: {len(product_data)}")
        _logger.info(f"Templates amount: {len(template_data)}")

        # Marcar tipo de producto
        for item in product_data:
            item["product_type"] = "product"

        for item in template_data:
            item["product_type"] = "template"

        # Crear conjunto de default_code de productos para evitar duplicados
        product_codes = {
            item.get("default_code")
            for item in product_data
            if item.get("default_code")
        }

        # Filtrar templates que no tengan el mismo default_code que los productos
        filtered_template_data = [
            item
            for item in template_data
            if item.get("default_code") not in product_codes
        ]

        _logger.info(f"Templates after deduplication: {len(filtered_template_data)}")

        # Combinar productos y templates filtrados
        all_products = product_data + filtered_template_data

        return all_products

    @api.model
    def get_product_by_sku(self, sku):
        return self.get_product(sku=sku)

    @api.model
    def get_product_by_name(self, name):
        return self.get_product(name=name)

    @api.model
    def get_product_by_id(self, id):
        return self.get_product(id=id)

    @api.model
    def get_sale_order(self, name=None, id=None):
        domain = [("amount_total", ">", 0)]
        if name:
            domain.append(("name", "ilike", name))
        elif id:
            domain.append(("id", "=", id))

        fields = [
            "id",
            "name",
            "partner_id",
            "date_order",
            "order_line",
            "state",
            "amount_total",
            "user_id",
            "company_id",
            "access_token",
        ]

        sale_order = self.env["sale.order"].sudo().search(domain, limit=1)
        if sale_order:
            result = sale_order.read(fields)[0]
            link = f"{self.env['ir.config_parameter'].sudo().get_param('web.base.url')}/my/orders/{sale_order.id}?access_token={sale_order.access_token}"
            result["link"] = link

            # Obtener detalles de productos
            result["products"] = []
            for line in sale_order.order_line:
                result["products"].append(
                    {
                        "product_id": line.product_id.id,
                        "product_name": line.product_id.name,
                        "sku": line.product_id.default_code,
                        "quantity": line.product_uom_qty,
                        "unit_price": line.price_unit,
                        "subtotal": line.price_subtotal,
                    }
                )

            return result

        _logger.warning(_("No se encontró el pedido %s"), name or id)
        return None

    @api.model
    def get_sale_order_by_name(self, name):
        _logger.info(f"Getting sale order {name}")
        return self.get_sale_order(name=name)

    @api.model
    def get_sale_order_by_id(self, id):
        _logger.info(f"Getting sale order {id}")
        return self.get_sale_order(id=id)

    @api.model
    def create_lead(self, partner, resume, email):
        try:
            lead_vals = {
                "stage_id": 1,
                "type": "opportunity",
                "name": f"WhatsApp - {partner['name']}",
                "email_from": email,
                "phone": partner["phone"],
                "description": resume,
                "partner_id": partner["id"],
            }
            lead = self.env["crm.lead"].sudo().create(lead_vals)
            if not lead:
                _logger.error("create_lead: La creación del lead falló. lead es None")
                return None

            lead_data = lead.read(["id", "name"])
            if not lead_data:
                _logger.error("create_lead: read() del lead falló")
                return None

            return lead_data[0]

        except Exception as e:
            _logger.error(f"create_lead: Excepción capturada: {e}")
            return None

    @api.model
    def create_partner(self, name, phone=None, email=None):
        if phone:
            phone = format_phone_number(phone)
            if not phone:
                return None, f"Número de teléfono inválido: {phone}"

            existing_partner = self.get_partner_by_phone(phone)
            if existing_partner:
                return existing_partner, "ALREADY"

        if email:
            existing_partner = self.get_partner_by_email(email)
            if existing_partner:
                return existing_partner, "ALREADY"

        partner_vals = {"name": name}
        if phone:
            partner_vals["phone"] = phone
        if email:
            partner_vals["email"] = email

        new_partner = self.env["res.partner"].sudo().create(partner_vals)
        return new_partner.read()[0], "CREATE"

    @api.model
    def create_sale_order(self, partner_id, order_line):
        order_vals = {
            "partner_id": partner_id,
            "order_line": [(0, 0, line) for line in order_line],
            "company_id": self.env.user.company_id.id,
            "access_token": secrets.token_urlsafe(32),
        }
        sale_order = self.env["sale.order"].sudo().create(order_vals)
        return sale_order.id

    @api.model
    def get_products_by_category(self, category_name):
        categories = (
            self.env["product.category"]
            .sudo()
            .search([("name", "ilike", category_name)])
        )
        product_data = []
        for category in categories:
            products = (
                self.env["product.product"]
                .sudo()
                .search([("categ_id", "=", category.id), ("active", "=", True)])
            )
            product_data.extend(
                products.read(["id", "name", "list_price", "qty_available"])
            )
        return product_data

    @api.model
    def get_category(self, id=None, name=None, parent_id=None):
        domain = []
        if id:
            domain.append(("id", "=", id))
        if name:
            domain.append(("name", "ilike", name))
        if parent_id:
            domain.append(("parent_id", "=", parent_id))

        categories = self.env["product.category"].sudo().search(domain)
        return categories.read(["id", "name", "parent_id", "product_count"])

    @api.model
    def get_all_categories(self):
        categories = self.env["product.category"].sudo().search([])
        return categories.read(["id", "name", "parent_id", "product_count"])

    @api.model
    def get_category_by_id(self, id):
        category = (
            self.env["product.category"].sudo().search([("id", "=", id)], limit=1)
        )
        return category.read(["id", "name", "parent_id", "product_count"])

    @api.model
    def get_categories_by_name(self, name):
        categories = (
            self.env["product.category"].sudo().search([("name", "ilike", name)])
        )
        return categories.read(["id", "name", "parent_id", "product_count"])

    @api.model
    def get_categories_children(self, parent_id):
        categories = (
            self.env["product.category"].sudo().search([("parent_id", "=", parent_id)])
        )
        return categories.read(["id", "name", "parent_id", "product_count"])

    @api.model
    def get_category_parent(self, child_id):
        category = self.env["product.category"].search([("id", "=", child_id)], limit=1)
        parent_category = category.parent_id
        return (
            parent_category.read(["id", "name", "parent_id", "product_count"])
            if parent_category
            else None
        )

    @api.model
    def sale_orders_by_user_id(self, partner_id):
        orders = self.env["sale.order"].sudo().search([("partner_id", "=", partner_id)])
        order_list = []
        for order in orders:
            link = f"{self.env['ir.config_parameter'].sudo().get_param('web.base.url')}/my/orders/{order.id}?access_token={order.access_token}"
            order_data = order.read(
                ["id", "name", "date_order", "state", "amount_total"]
            )[0]
            order_data["link"] = link
            order_list.append(order_data)

        return order_list

    # Crear línea de producto para un pedido
    @api.model
    def create_order_line(self, products):
        order_lines = []
        for product_info in products:
            product = (
                self.env["product.product"]
                .sudo()
                .search([("default_code", "=", product_info["default_code"])], limit=1)
            )
            if not product:
                _logger.warning(
                    _("Producto con SKU %s no existe"), product_info["default_code"]
                )
                continue

            if product.qty_available < product_info["uom_qty"]:
                _logger.warning(
                    _("Producto %s con SKU %s está agotado"),
                    product.name,
                    product.default_code,
                )
                continue

            order_line_data = {
                "product_id": product.id,
                "product_uom_qty": product_info["uom_qty"],
                "price_unit": product.list_price * product_info["uom_qty"],
            }
            order_lines.append(order_line_data)
        return order_lines

    # Obtener IDs de categorías hijas para una categoría dada
    @api.model
    def get_children_ids(self, category_id):
        category = self.env["product.category"].sudo().browse(category_id)
        if not category.exists():
            _logger.warning(_("Categoría con ID %s no existe"), category_id)
            return []
        all_children = category.child_id
        return all_children.mapped("id")

    @api.model
    def get_products_by_category_id(self, category_id):
        category = self.env["product.category"].sudo().browse(category_id)
        if not category.exists():
            _logger.warning("Categoría con ID %s no existe", category_id)
            return []

        child_categories = category.child_id
        all_categories = child_categories + category

        domain = [("categ_id", "in", all_categories.ids), ("active", "=", True)]
        products = self.env["product.product"].sudo().search(domain)
        return products.read(["id", "name", "list_price", "qty_available"])

    @api.model
    def get_products_by_category_name(self, category_name):
        categories = (
            self.env["product.category"]
            .sudo()
            .search([("name", "ilike", category_name)])
        )
        products_by_category = {}

        for category in categories:
            products = self.get_products_by_category_id(category.id)
            products_by_category[category.name] = products

        return products_by_category

    # ======= CALENDAR EVENT METHODS =======

    @api.model
    def create_calendar_event(
        self,
        name,
        start_datetime,
        end_datetime=None,
        description=None,
        partner_ids=None,
        location=None,
        allday=False,
        duration=1.0,
    ):
        """
        Crear un evento de calendario.

        Args:
            name (str): Nombre del evento
            start_datetime (str): Fecha y hora de inicio
            end_datetime (str, optional): Fecha y hora de fin
            description (str, optional): Descripción del evento
            partner_ids (list, optional): Lista de IDs de contactos
            location (str, optional): Ubicación del evento
            allday (bool, optional): Evento de todo el día
            duration (float, optional): Duración en horas
        """
        calendar_model = self.env["calendar.event"]
        return calendar_model.create_calendar_event(
            name=name,
            start_datetime=start_datetime,
            end_datetime=end_datetime,
            description=description,
            partner_ids=partner_ids,
            location=location,
            allday=allday,
            duration=duration,
        )

    @api.model
    def get_calendar_events(
        self,
        start_date=None,
        end_date=None,
        partner_id=None,
        limit=20,
        search_term=None,
    ):
        """
        Consultar eventos de calendario.

        Args:
            start_date (str, optional): Fecha de inicio del rango
            end_date (str, optional): Fecha de fin del rango
            partner_id (int, optional): ID del contacto
            limit (int, optional): Límite de resultados
            search_term (str, optional): Término de búsqueda
        """
        calendar_model = self.env["calendar.event"]
        return calendar_model.get_calendar_events(
            start_date=start_date,
            end_date=end_date,
            partner_id=partner_id,
            limit=limit,
            search_term=search_term,
        )

    @api.model
    def get_upcoming_events(self, days_ahead=7, limit=10):
        """
        Obtener eventos próximos.

        Args:
            days_ahead (int, optional): Días hacia adelante
            limit (int, optional): Límite de resultados
        """
        calendar_model = self.env["calendar.event"]
        return calendar_model.get_upcoming_events(days_ahead=days_ahead, limit=limit)

    @api.model
    def update_calendar_event(self, event_id, **kwargs):
        """
        Actualizar un evento de calendario.

        Args:
            event_id (int): ID del evento
            **kwargs: Campos a actualizar
        """
        calendar_model = self.env["calendar.event"]
        return calendar_model.update_calendar_event(event_id, **kwargs)

    @api.model
    def delete_calendar_event(self, event_id):
        """
        Eliminar un evento de calendario.

        Args:
            event_id (int): ID del evento
        """
        calendar_model = self.env["calendar.event"]
        return calendar_model.delete_calendar_event(event_id)
